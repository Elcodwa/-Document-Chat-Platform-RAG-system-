"""
Turns an uploaded file's raw bytes into plain text (plus, where the format
supports it, a page number for each piece of text - used later so
citations can point to "page 4 of contract.pdf" rather than just the
file name).
"""
from dataclasses import dataclass, field

import pandas as pd
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from app.exceptions import FileProcessingError

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md"}


@dataclass
class ExtractedPage:
    page_number: int | None
    text: str


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage] = field(default_factory=list)
    page_count: int | None = None

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())


def extract_text(file_path: str, extension: str) -> ExtractedDocument:
    extension = extension.lower()
    try:
        if extension == ".pdf":
            return _extract_pdf(file_path)
        if extension == ".docx":
            return _extract_docx(file_path)
        if extension in (".xlsx", ".xls"):
            return _extract_excel(file_path)
        if extension == ".csv":
            return _extract_csv(file_path)
        if extension in (".txt", ".md"):
            return _extract_plain_text(file_path)
    except FileProcessingError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise FileProcessingError(f"Could not read file contents ({extension}).") from exc

    raise FileProcessingError(f"Unsupported file type: {extension}")


def _extract_pdf(file_path: str) -> ExtractedDocument:
    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(ExtractedPage(page_number=i, text=text))
    return ExtractedDocument(pages=pages, page_count=len(reader.pages))


def _extract_docx(file_path: str) -> ExtractedDocument:
    doc = DocxDocument(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    # DOCX has no fixed pagination in the file format itself, so we treat
    # the whole document as one logical "page" for citation purposes.
    return ExtractedDocument(pages=[ExtractedPage(page_number=None, text=text)], page_count=None)


def _extract_excel(file_path: str) -> ExtractedDocument:
    workbook = load_workbook(file_path, data_only=True, read_only=True)
    pages = []
    for sheet_index, sheet_name in enumerate(workbook.sheetnames, start=1):
        sheet = workbook[sheet_name]
        rows_text = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(v) for v in row if v is not None]
            if values:
                rows_text.append(" | ".join(values))
        if rows_text:
            pages.append(ExtractedPage(page_number=sheet_index, text=f"Sheet: {sheet_name}\n" + "\n".join(rows_text)))
    return ExtractedDocument(pages=pages, page_count=len(workbook.sheetnames))


def _extract_csv(file_path: str) -> ExtractedDocument:
    df = pd.read_csv(file_path)
    text = df.to_string(index=False)
    return ExtractedDocument(pages=[ExtractedPage(page_number=None, text=text)], page_count=None)


def _extract_plain_text(file_path: str) -> ExtractedDocument:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return ExtractedDocument(pages=[ExtractedPage(page_number=None, text=text)], page_count=None)
